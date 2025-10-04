import os
import io
import tempfile
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from openai import OpenAI
import markdown
from bs4 import BeautifulSoup

# Load environment
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DEFAULT_MODEL = os.getenv("DEFAULT_MODEL", "gpt-4o")

if not OPENAI_API_KEY:
    st.error("‚ö†Ô∏è No API key found. Please set OPENAI_API_KEY in your .env file.")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# --- Streamlit UI ---
st.title("üñºÔ∏è Image ‚Üí AI Report Generator (GPT-4o)")
st.write("Upload an image (PNG/JPG). The AI will analyze it and generate a structured report.")

uploaded_file = st.file_uploader("Choose an image", type=["png", "jpg", "jpeg"])

model_choice = st.selectbox("Choose model", [DEFAULT_MODEL, "gpt-4o", "gpt-4o-mini"], index=0)

user_prompt = st.text_area(
    "Write what you want the AI to do with the image:",
    value="""Extract the data from this image and generate a Galvanic Skin Response (GSR) Health Report with the following. Identify baseline levels and changes in conductance.Point out peaks or significant fluctuations. Correlate variations with possible stress, arousal, or relaxation responses. Comment on symmetry, stability, and response intensity.""",
    height=200
)

# --- Helpers ---
def generate_report_from_image(image_bytes, model, prompt_text):
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a naturopathic practitioner specializing in biometric data interpretation. Your task is to generate clear, detailed, comprehensive and structured reports based on Galvanic Skin Response (GSR) readings."},
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text},
                    {"type": "image_url", "image_url": {"url": "data:image/png;base64," + image_bytes}},
                ],
            },
        ],
        max_tokens=1200,
    )
    return resp.choices[0].message.content

def create_docx(report_text, title="Generated Report"):
    html = markdown.markdown(report_text)
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    doc.add_heading(title, level=1)
    
    for element in soup.descendants:
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# --- Main flow ---
if uploaded_file:
    if st.button("Generate Report"):
        with st.spinner("Analyzing image with GPT-4o‚Ä¶"):
            import base64
            b64_image = base64.b64encode(uploaded_file.read()).decode("utf-8")

            try:
                report = generate_report_from_image(b64_image, model_choice, user_prompt)

                st.subheader("Generated Report")
                st.text_area("Report (editable)", value=report, height=400)

                path = create_docx(report)
                with open(path, "rb") as f:
                    st.download_button("Download as Word (.docx)", f, file_name="Image_Report.docx")
            except Exception as e:
                st.error(f"Error generating report: {e}")